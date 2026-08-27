from __future__ import annotations
import os
import re
from pathlib import Path
from typing import Any, List, Optional, Tuple
import docx
from docx import Document
from core.model import WordMatch, WordReplaceRule
from core.redact.mask import DEFAULT_MASK_CHAR, mask_placeholder
from core.detector.rule_engine import RuleEngine
from core.word.replacer import apply_range_to_runs, replace_paragraph_text
from core.word.rules_loader import load_word_replace_rules, merge_word_replace_rules, load_all_word_detection_rules

_PII_REPLACEMENT_HINTS = (
    ('护照', DEFAULT_MASK_CHAR),
    ('身份证', DEFAULT_MASK_CHAR),
    ('手机', DEFAULT_MASK_CHAR),
    ('邮箱', DEFAULT_MASK_CHAR),
    ('银行卡', DEFAULT_MASK_CHAR),
    ('电话', DEFAULT_MASK_CHAR),
    ('座机', DEFAULT_MASK_CHAR),
    ('港澳', DEFAULT_MASK_CHAR),
    ('国际', DEFAULT_MASK_CHAR),
)


class WordPipeline:
    def __init__(
        self,
        rule_engine: Optional[RuleEngine] = None,
        rules: Optional[List[WordReplaceRule]] = None,
        *,
        config_rules: Optional[List[WordReplaceRule]] = None,
    ):
        self.rule_engine = rule_engine
        config = list(config_rules) if config_rules is not None else load_word_replace_rules()
        self.rules = merge_word_replace_rules(config, rules)

    @classmethod
    def from_document_rules(
        cls,
        rule_engine: Optional[RuleEngine] = None,
        extra_rules: Optional[List[WordReplaceRule]] = None,
    ) -> WordPipeline:
        engine = rule_engine or RuleEngine.load_document()
        merged_rules = load_all_word_detection_rules(extra_rules=extra_rules)
        return cls(rule_engine=engine, rules=merged_rules, config_rules=[])

    def _engine_replacement(self, rule_label: str, matched_text: str) -> str:
        label = rule_label.strip('[]')
        configured = DEFAULT_MASK_CHAR
        for hint, replacement in _PII_REPLACEMENT_HINTS:
            if hint in label:
                configured = replacement
                break
        return mask_placeholder(matched_text, configured)

    @staticmethod
    def _resolve_replacement(matched_text: str, configured: str) -> str:
        return mask_placeholder(matched_text, configured)

    @staticmethod
    def _match_to_dict(match: WordMatch) -> dict[str, Any]:
        return {
            'start': match.start,
            'end': match.end,
            'text': match.text,
            'replacement': match.replacement,
            'type': match.source,
            'mode': match.mode,
        }

    @staticmethod
    def _iter_document_paragraphs(doc: Document):
        for p_idx, para in enumerate(doc.paragraphs):
            yield 'body', p_idx, para

        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, para in enumerate(cell.paragraphs):
                        yield f'table:{t_idx}:{r_idx}:{c_idx}', p_idx, para

        for s_idx, section in enumerate(doc.sections):
            if section.header:
                for p_idx, para in enumerate(section.header.paragraphs):
                    yield f'header:{s_idx}', p_idx, para
            if section.footer:
                for p_idx, para in enumerate(section.footer.paragraphs):
                    yield f'footer:{s_idx}', p_idx, para

    def scan_document(self, doc: Document) -> dict[str, Any]:
        paragraphs_data: list[dict[str, Any]] = []
        tables_data: list[dict[str, Any]] = []
        total_matches = 0
        matches_summary: dict[str, int] = {}

        table_buckets: dict[int, dict[str, Any]] = {}

        for location, p_idx, para in self._iter_document_paragraphs(doc):
            text = para.text
            if not text.strip():
                continue
            matches = self.find_matches_in_text(text)
            if matches:
                total_matches += len(matches)
                for match in matches:
                    matches_summary[match.text] = matches_summary.get(match.text, 0) + 1

            block = {
                'index': p_idx,
                'location': location,
                'text': text,
                'matches': [self._match_to_dict(m) for m in matches],
            }

            if location == 'body':
                paragraphs_data.append(block)
            elif location.startswith('table:'):
                parts = location.split(':')
                table_index = int(parts[1])
                row_index = int(parts[2])
                col_index = int(parts[3])
                bucket = table_buckets.setdefault(
                    table_index,
                    {'table_index': table_index, 'rows': {}},
                )
                rows = bucket['rows']
                row = rows.setdefault(row_index, [])
                row.append(
                    {
                        'col_index': col_index,
                        'text': text,
                        'matches_count': len(matches),
                        'matches': block['matches'],
                    }
                )
            else:
                paragraphs_data.append(block)

        for table_index in sorted(table_buckets.keys()):
            bucket = table_buckets[table_index]
            rows_data = []
            for row_index in sorted(bucket['rows'].keys()):
                cells = sorted(bucket['rows'][row_index], key=lambda c: c['col_index'])
                rows_data.append(cells)
            tables_data.append({'table_index': table_index, 'rows': rows_data})

        return {
            'paragraphs': paragraphs_data,
            'tables': tables_data,
            'total_matches': total_matches,
            'matches_summary': matches_summary,
        }

    def find_matches_in_text(self, text: str) -> List[WordMatch]:
        if not text:
            return []
        matches: List[WordMatch] = []
        # 1. Check custom WordReplaceRules
        for idx, rule in enumerate(self.rules):
            if not rule.enabled or not rule.find:
                continue
            if rule.mode == 'regex':
                try:
                    for m in re.finditer(rule.find, text, re.IGNORECASE):
                        matched = m.group(0)
                        matches.append(WordMatch(
                            start=m.start(),
                            end=m.end(),
                            text=matched,
                            replacement=self._resolve_replacement(matched, rule.replace),
                            mode='regex',
                            source='rule',
                            rule_index=idx
                        ))
                except Exception:
                    pass
            else:
                start = 0
                find_str = rule.find
                while True:
                    pos = text.find(find_str, start)
                    if pos == -1:
                        break
                    matches.append(WordMatch(
                        start=pos,
                        end=pos + len(find_str),
                        text=find_str,
                        replacement=self._resolve_replacement(find_str, rule.replace),
                        mode='exact',
                        source='rule',
                        rule_index=idx
                    ))
                    start = pos + max(1, len(find_str))

        # 2. Check rule_engine terms / PII
        if self.rule_engine:
            spans = self.rule_engine.find_spans(text)
            for s, e, term in spans:
                # avoid overlapping with existing custom rule matches
                if not any(not (e <= m.start or s >= m.end) for m in matches):
                    matches.append(WordMatch(
                        start=s,
                        end=e,
                        text=text[s:e],
                        replacement=self._engine_replacement(term, text[s:e]),
                        mode='engine',
                        source='engine',
                        rule_index=-1
                    ))

        # Sort and deduplicate overlapping matches (prefer longer match)
        matches.sort(key=lambda m: (m.start, -(m.end - m.start)))
        filtered = []
        last_end = -1
        for m in matches:
            if m.start >= last_end:
                filtered.append(m)
                last_end = m.end
        return filtered

    def process_document(self, input_path: str | Path, output_path: str | Path) -> dict[str, Any]:
        in_p = Path(input_path)
        out_p = Path(output_path)
        
        doc = Document(str(in_p))
        total_replacements = 0
        hit_records = []

        # Process body paragraphs
        for para in doc.paragraphs:
            matches = self.find_matches_in_text(para.text)
            if matches:
                hit_records.extend([m.text for m in matches])
                total_replacements += replace_paragraph_text(para, matches)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        matches = self.find_matches_in_text(para.text)
                        if matches:
                            hit_records.extend([m.text for m in matches])
                            total_replacements += replace_paragraph_text(para, matches)

        # Process headers and footers
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    matches = self.find_matches_in_text(para.text)
                    if matches:
                        hit_records.extend([m.text for m in matches])
                        total_replacements += replace_paragraph_text(para, matches)
            if section.footer:
                for para in section.footer.paragraphs:
                    matches = self.find_matches_in_text(para.text)
                    if matches:
                        hit_records.extend([m.text for m in matches])
                        total_replacements += replace_paragraph_text(para, matches)

        out_p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_p))
        return {
            'input_path': str(in_p),
            'output_path': str(out_p),
            'total_replacements': total_replacements,
            'matches_count': total_replacements,
            'hit_terms': list(set(hit_records))
        }
