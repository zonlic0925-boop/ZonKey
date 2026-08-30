"""转换工具桥测试（backend_convert_tools）。

合成样本全部现场生成（reportlab / python-docx / openpyxl），
Office COM 用例依赖本机 Word/Excel，缺失时自动跳过。
产物校验后立即清理，不污染 output/。
"""

from __future__ import annotations

import io
import time
from pathlib import Path

import pytest
from fastapi import FastAPI
from fastapi.testclient import TestClient

import backend_convert_tools as convert_tools


def _make_client() -> TestClient:
    app = FastAPI()
    app.include_router(convert_tools.router)
    return TestClient(app)


def _wait_job(client: TestClient, job_id: str, timeout: float = 120) -> dict:
    deadline = time.time() + timeout
    last = {}
    while time.time() < deadline:
        response = client.get(f"/api/convert/jobs/{job_id}")
        assert response.status_code == 200
        last = response.json()
        if last["status"] in ("done", "error"):
            return last
        time.sleep(0.2)
    raise AssertionError(f"任务 {job_id} 超时未完成: {last}")


# ---------------------------------------------------------------------------
# 合成样本
# ---------------------------------------------------------------------------

def _sample_pdf_bytes() -> bytes:
    pytest.importorskip("reportlab")
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    except Exception:
        pass  # CID 字体注册失败时退化为纯英文样本

    styles = getSampleStyleSheet()
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, title="ZonScale convert sample")
    story = [
        Paragraph("ZonScale Convert Marker 42", styles["Title"]),
        Spacer(1, 12),
        Paragraph("This is the first body paragraph for conversion testing.", styles["BodyText"]),
        Paragraph("中文段落：图纸脱敏工作台转换引擎测试。", styles["BodyText"]),
        Spacer(1, 12),
        Table(
            [["Part", "Qty", "Price"], ["Widget-A", "42.5", "9.99"], ["Gadget-B", "7", "199.00"]],
            colWidths=[180, 90, 90],
            style=TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ]),
        ),
        PageBreak(),
        Paragraph("Second page heading", styles["Heading1"]),
        Paragraph("Second page body content for page splitting check.", styles["BodyText"]),
    ]
    doc.build(story)
    return buffer.getvalue()


def _sample_docx_bytes() -> bytes:
    pytest.importorskip("docx")
    from docx import Document

    doc = Document()
    doc.add_heading("Office COM Marker 7", level=0)
    doc.add_paragraph("Word to PDF conversion body text.")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _sample_xlsx_bytes() -> bytes:
    pytest.importorskip("openpyxl")
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet["A1"] = "Excel COM Marker 9"
    sheet["B1"] = 3.14
    sheet.column_dimensions["A"].width = 25  # 拉宽避免 PDF 导出时溢出文本被右侧单元格截断
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


@pytest.fixture(scope="module")
def sample_pdf() -> bytes:
    return _sample_pdf_bytes()


@pytest.fixture(scope="module")
def has_office_com() -> bool:
    return convert_tools._probe_win32()


# ---------------------------------------------------------------------------
# 基础行为
# ---------------------------------------------------------------------------

def test_capability_flags():
    client = _make_client()
    response = client.get("/api/convert/capability")
    assert response.status_code == 200
    data = response.json()
    assert isinstance(data["engines"], dict) and data["engines"]
    assert isinstance(data["word_com"], bool)
    assert isinstance(data["excel_com"], bool)


def test_unknown_job_404():
    client = _make_client()
    assert client.get("/api/convert/jobs/does-not-exist").status_code == 404


def test_pdf_upload_rejects_non_pdf():
    client = _make_client()
    response = client.post(
        "/api/convert/pdf-to-word",
        files={"file": ("notes.pdf", io.BytesIO(b"plain text, not a pdf"), "text/plain")},
    )
    assert response.status_code == 400


def test_pdf_upload_rejects_too_small():
    client = _make_client()
    response = client.post(
        "/api/convert/pdf-to-word",
        files={"file": ("tiny.pdf", io.BytesIO(b"%PDF-1.7 broken"), "application/pdf")},
    )
    assert response.status_code == 400


def test_office_upload_rejects_pptx_with_hint():
    client = _make_client()
    response = client.post(
        "/api/convert/office-to-pdf",
        files={"file": ("deck.pptx", io.BytesIO(b"PK\x03\x04"), "application/octet-stream")},
    )
    assert response.status_code == 400
    assert "PPT 工坊" in response.json()["detail"]


# ---------------------------------------------------------------------------
# PDF→Word
# ---------------------------------------------------------------------------

def test_pdf_to_word_end_to_end(sample_pdf: bytes):
    client = _make_client()
    response = client.post(
        "/api/convert/pdf-to-word",
        files={"file": ("zonscale_sample.pdf", io.BytesIO(sample_pdf), "application/pdf")},
    )
    assert response.status_code == 200, response.text
    data = _wait_job(client, response.json()["job_id"])
    assert data["status"] == "done", data

    out_path = Path(data["outputs"][0]["dir"]) / data["outputs"][0]["name"]
    try:
        assert out_path.exists() and out_path.suffix == ".docx"
        assert data["page_count"] == 2
        assert data["table_count"] >= 1
        from docx import Document

        doc = Document(str(out_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "ZonScale Convert Marker 42" in full_text
        table_text = "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        assert "Widget-A" in table_text
    finally:
        out_path.unlink(missing_ok=True)


# ---------------------------------------------------------------------------
# PDF→Excel
# ---------------------------------------------------------------------------

def test_pdf_to_excel_end_to_end(sample_pdf: bytes):
    client = _make_client()
    response = client.post(
        "/api/convert/pdf-to-excel",
        files={"file": ("zonscale_sample.pdf", io.BytesIO(sample_pdf), "application/pdf")},
    )
    assert response.status_code == 200, response.text
    data = _wait_job(client, response.json()["job_id"])
    assert data["status"] == "done", data

    out_path = Path(data["outputs"][0]["dir"]) / data["outputs"][0]["name"]
    try:
        assert out_path.exists() and out_path.suffix == ".xlsx"
        assert data["table_count"] >= 1
        from openpyxl import load_workbook

        workbook = load_workbook(str(out_path))
        assert any(ws.title.startswith("Page") for ws in workbook.worksheets)
        values = [
            cell.value
            for ws in workbook.worksheets
            for row in ws.iter_rows()
            for cell in row
            if cell.value is not None
        ]
        assert "Widget-A" in values
        assert 42.5 in values  # 纯数字单元格转为数值类型
    finally:
        out_path.unlink(missing_ok=True)


def test_pdf_to_excel_no_tables_errors(sample_pdf: bytes):
    pytest.importorskip("reportlab")
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate

    styles = getSampleStyleSheet()
    buffer = io.BytesIO()
    SimpleDocTemplate(buffer, pagesize=A4).build(
        [Paragraph("Only prose, no ruled tables here at all.", styles["BodyText"])]
    )

    client = _make_client()
    response = client.post(
        "/api/convert/pdf-to-excel",
        files={"file": ("prose_only.pdf", io.BytesIO(buffer.getvalue()), "application/pdf")},
    )
    assert response.status_code == 200
    data = _wait_job(client, response.json()["job_id"])
    assert data["status"] == "error"
    assert "表格" in (data.get("error") or "")


# ---------------------------------------------------------------------------
# PDF→PPT
# ---------------------------------------------------------------------------

def test_pdf_to_ppt_end_to_end(sample_pdf: bytes):
    pytest.importorskip("pypdfium2")
    pytest.importorskip("pptx")
    client = _make_client()
    response = client.post(
        "/api/convert/pdf-to-ppt",
        files={"file": ("zonscale_sample.pdf", io.BytesIO(sample_pdf), "application/pdf")},
        data={"dpi": "110"},
    )
    assert response.status_code == 200, response.text
    data = _wait_job(client, response.json()["job_id"])
    assert data["status"] == "done", data
    assert data["page_count"] == 2
    assert data["dpi"] == 110

    out_path = Path(data["outputs"][0]["dir"]) / data["outputs"][0]["name"]
    try:
        assert out_path.exists() and out_path.suffix == ".pptx"
        from pptx import Presentation

        prs = Presentation(str(out_path))
        assert len(prs.slides) == 2
        expected_width = 595.2755905511812 * 12700  # A4 宽度 pt → EMU
        assert abs(prs.slide_width - expected_width) / expected_width < 0.02
    finally:
        out_path.unlink(missing_ok=True)


# ---------------------------------------------------------------------------
# Office→PDF（COM）
# ---------------------------------------------------------------------------

@pytest.mark.skipif(not convert_tools._probe_win32(), reason="本机无 pywin32/Office COM")
def test_office_to_pdf_word_end_to_end(has_office_com: bool):
    client = _make_client()
    response = client.post(
        "/api/convert/office-to-pdf",
        files={"file": ("zonscale_word_sample.docx", io.BytesIO(_sample_docx_bytes()), "application/octet-stream")},
    )
    assert response.status_code == 200, response.text
    data = _wait_job(client, response.json()["job_id"])
    assert data["status"] == "done", data
    assert data["office_app"] == "word"

    out_path = Path(data["outputs"][0]["dir"]) / data["outputs"][0]["name"]
    try:
        assert out_path.exists() and out_path.suffix == ".pdf"
        from core.pdfio import open_doc

        with open_doc(str(out_path)) as pdf:
            assert pdf.page_count >= 1
            assert "Office COM Marker 7" in pdf.page(0).page_text()
    finally:
        out_path.unlink(missing_ok=True)


@pytest.mark.skipif(not convert_tools._probe_win32(), reason="本机无 pywin32/Office COM")
def test_office_to_pdf_excel_end_to_end(has_office_com: bool):
    client = _make_client()
    response = client.post(
        "/api/convert/office-to-pdf",
        files={"file": ("zonscale_excel_sample.xlsx", io.BytesIO(_sample_xlsx_bytes()), "application/octet-stream")},
    )
    assert response.status_code == 200, response.text
    data = _wait_job(client, response.json()["job_id"])
    assert data["status"] == "done", data
    assert data["office_app"] == "excel"

    out_path = Path(data["outputs"][0]["dir"]) / data["outputs"][0]["name"]
    try:
        assert out_path.exists() and out_path.suffix == ".pdf"
        from core.pdfio import open_doc

        with open_doc(str(out_path)) as pdf:
            assert pdf.page_count >= 1
            assert "Excel COM Marker 9" in pdf.page(0).page_text()
    finally:
        out_path.unlink(missing_ok=True)


# ---------------------------------------------------------------------------
# PDF 修复（pikepdf）
# ---------------------------------------------------------------------------

def test_repair_valid_pdf(sample_pdf: bytes):
    client = _make_client()
    response = client.post(
        "/api/convert/repair",
        files={"file": ("zonscale_sample.pdf", io.BytesIO(sample_pdf), "application/pdf")},
    )
    assert response.status_code == 200, response.text
    data = response.json()
    assert data["status"] == "success"
    assert data["page_count"] == 2

    out_path = Path(data["output_dir"]) / data["download_name"]
    try:
        import pikepdf

        with pikepdf.open(str(out_path)) as pdf:
            assert len(pdf.pages) == 2
    finally:
        out_path.unlink(missing_ok=True)


def test_repair_unrecoverable_garbage():
    garbage = b"%PDF-1.7\n" + b"\x00" * 128 + b"trailer\n%%EOF\n"
    client = _make_client()
    response = client.post(
        "/api/convert/repair",
        files={"file": ("broken.pdf", io.BytesIO(garbage), "application/pdf")},
    )
    assert response.status_code == 422
    assert "修复" in response.json()["detail"]


# ---------------------------------------------------------------------------
# P2 端点补齐：compress-deep / html-to-pdf / ocr-export / office 兜底链
# ---------------------------------------------------------------------------

def _image_heavy_pdf_bytes() -> bytes:
    """大图 PDF（噪声图），保证栅格化重编码有压缩收益。"""
    import numpy as np
    from PIL import Image as PILImage
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen.canvas import Canvas

    buffer = io.BytesIO()
    canvas = Canvas(buffer, pagesize=(400, 400))
    noise = (np.random.rand(600, 800, 3) * 255).astype("uint8")
    canvas.drawImage(ImageReader(PILImage.fromarray(noise)), 0, 0, width=400, height=400)
    canvas.showPage()
    canvas.save()
    return buffer.getvalue()


def test_compress_deep_end_to_end():
    client = _make_client()
    original = _image_heavy_pdf_bytes()
    response = client.post(
        "/api/convert/compress-deep",
        files={"file": ("heavy.pdf", io.BytesIO(original), "application/pdf")},
        data={"dpi": "120", "quality": "60"},
    )
    assert response.status_code == 200, response.text
    data = _wait_job(client, response.json()["job_id"])
    assert data["status"] == "done", data
    assert data["original_bytes"] == len(original)
    assert data["compressed_bytes"] < len(original), "栅格化重编码应缩小图片型 PDF"
    assert "栅格化" in data["note"]  # 诚实标注：无可编辑文本层
    out_path = Path(data["outputs"][0]["dir"]) / data["outputs"][0]["name"]
    try:
        import pypdfium2 as pdfium

        doc = pdfium.PdfDocument(str(out_path))
        try:
            assert len(doc) == 1
        finally:
            doc.close()
    finally:
        out_path.unlink(missing_ok=True)


def test_html_to_pdf_markdown():
    client = _make_client()
    content = "# 标题一\n\n这是**粗体**段落。\n\n- 项目甲\n- 项目乙\n\n1. 第一\n2. 第二\n\n> 引用块\n\n---\n\n正文结尾。"
    response = client.post(
        "/api/convert/html-to-pdf",
        data={"content": content, "title": "测试文档"},
    )
    assert response.status_code == 200, response.text
    assert response.json()["input_format"] == "markdown"
    out_path = Path(response.json()["output_dir"]) / response.json()["download_name"]
    try:
        import pdfplumber

        with pdfplumber.open(str(out_path)) as pdf:
            text = pdf.pages[0].extract_text() or ""
        assert "标题一" in text
        assert "粗体" in text
        assert "项目甲" in text
        assert "第一" in text
    finally:
        out_path.unlink(missing_ok=True)


def test_html_to_pdf_html_subset_with_table():
    client = _make_client()
    html = (
        "<h1>Report</h1><p>Hello <b>world</b>.</p>"
        "<table><tr><th>A</th><th>B</th></tr><tr><td>x1</td><td>x2</td></tr></table>"
        "<ul><li>item one</li><li>item two</li></ul>"
    )
    response = client.post("/api/convert/html-to-pdf", data={"content": html})
    assert response.status_code == 200, response.text
    assert response.json()["input_format"] == "html"
    out_path = Path(response.json()["output_dir"]) / response.json()["download_name"]
    try:
        import pdfplumber

        with pdfplumber.open(str(out_path)) as pdf:
            text = pdf.pages[0].extract_text() or ""
        assert "Report" in text and "Hello world" in text
        assert "A" in text and "x1" in text and "x2" in text
        assert "item one" in text
    finally:
        out_path.unlink(missing_ok=True)


def test_html_to_pdf_empty_rejected():
    client = _make_client()
    response = client.post("/api/convert/html-to-pdf", data={"content": "   "})
    assert response.status_code == 400


def test_ocr_export_txt():
    client = _make_client()
    pytest.importorskip("rapidocr_onnxruntime")
    buffer = io.BytesIO()
    from reportlab.pdfgen.canvas import Canvas

    canvas = Canvas(buffer, pagesize=(400, 200))
    canvas.setFont("Helvetica", 40)
    canvas.drawString(40, 120, "ZONSCALE 2026")
    canvas.showPage()
    canvas.save()
    buffer.seek(0)
    response = client.post(
        "/api/convert/ocr-export",
        files={"file": ("sample.pdf", buffer, "application/pdf")},
        data={"output": "txt", "dpi": "220"},
    )
    assert response.status_code == 200, response.text
    data = _wait_job(client, response.json()["job_id"], timeout=180)
    assert data["status"] == "done", data
    assert data["char_count"] > 0
    out_path = Path(data["outputs"][0]["dir"]) / data["outputs"][0]["name"]
    try:
        text = out_path.read_text(encoding="utf-8")
        assert "ZONSCALE" in text.upper().replace(" ", "") or "2026" in text
    finally:
        out_path.unlink(missing_ok=True)


def test_ocr_export_sandwich_pdf_searchable():
    client = _make_client()
    pytest.importorskip("rapidocr_onnxruntime")
    buffer = io.BytesIO()
    from reportlab.pdfgen.canvas import Canvas

    canvas = Canvas(buffer, pagesize=(400, 200))
    canvas.setFont("Helvetica", 40)
    canvas.drawString(40, 120, "ZONSCALE 2026")
    canvas.showPage()
    canvas.save()
    buffer.seek(0)
    response = client.post(
        "/api/convert/ocr-export",
        files={"file": ("sample.pdf", buffer, "application/pdf")},
        data={"output": "pdf", "dpi": "220"},
    )
    assert response.status_code == 200, response.text
    data = _wait_job(client, response.json()["job_id"], timeout=180)
    assert data["status"] == "done", data
    out_path = Path(data["outputs"][0]["dir"]) / data["outputs"][0]["name"]
    try:
        import pdfplumber

        with pdfplumber.open(str(out_path)) as pdf:
            text = pdf.pages[0].extract_text() or ""
        # 夹心 PDF：隐形文字层可搜索（OCR 结果）
        assert "2026" in text or "ZONSCALE" in text.upper()
    finally:
        out_path.unlink(missing_ok=True)


def test_office_word_fallback_chain():
    """Word 兜底链：mammoth docx→HTML→platypus（无 COM 依赖，CJK 可用）。"""
    pytest.importorskip("mammoth")
    from docx import Document

    doc = Document()
    doc.add_heading("兜底标题", level=1)
    doc.add_paragraph("这是兜底转换的正文段落。")
    source = Path("temp_bridge_files") / "_fb_test_fallback.docx"
    source.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(source))
    out_path = Path("output") / "_fb_test_fallback.pdf"
    try:
        job_id = convert_tools._new_job("office-fallback-test")
        convert_tools._word_to_pdf_fallback(source, out_path, job_id)
        import pdfplumber

        with pdfplumber.open(str(out_path)) as pdf:
            text = pdf.pages[0].extract_text() or ""
        assert "兜底标题" in text
        assert "正文段落" in text
    finally:
        source.unlink(missing_ok=True)
        out_path.unlink(missing_ok=True)


def test_office_excel_fallback_chain():
    """Excel 兜底链：openpyxl→HTML 表格→platypus。"""
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "数据"
    sheet["A1"] = "名称"
    sheet["B1"] = "数量"
    sheet["A2"] = "零件甲"
    sheet["B2"] = 42
    source = Path("temp_bridge_files") / "_fb_test_excel.xlsx"
    source.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(str(source))
    out_path = Path("output") / "_fb_test_excel.pdf"
    try:
        job_id = convert_tools._new_job("office-fallback-test")
        convert_tools._excel_to_pdf_fallback(source, out_path, job_id)
        import pdfplumber

        with pdfplumber.open(str(out_path)) as pdf:
            text = pdf.pages[0].extract_text() or ""
        assert "数据" in text
        assert "零件甲" in text and "42" in text
    finally:
        source.unlink(missing_ok=True)
        out_path.unlink(missing_ok=True)


def test_capability_reports_new_engines():
    client = _make_client()
    response = client.get("/api/convert/capability")
    assert response.status_code == 200
    data = response.json()
    assert "rapidocr" in data
    assert "word_fallback" in data
