import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from core.model import WordReplaceRule
from core.word.pipeline import WordPipeline
from core.detector.rule_engine import RuleEngine

def test_word_pipeline_exact_replace(tmp_path):
    doc_path = tmp_path / 'test_sample.docx'
    out_path = tmp_path / 'test_sample_out.docx'
    
    doc = Document()
    p1 = doc.add_paragraph()
    r1 = p1.add_run('甲方：北京科技有限公司，')
    r1.bold = True
    r2 = p1.add_run('乙方：张三先生。')
    r2.font.color.rgb = RGBColor(255, 0, 0)
    
    table = doc.add_table(rows=1, cols=2)
    cell_p = table.cell(0, 0).paragraphs[0]
    cell_p.add_run('联系电话：13800138000')
    
    doc.save(str(doc_path))
    
    # Run WordPipeline with custom rule and PII regex
    rules = [
        WordReplaceRule(find='北京科技有限公司', replace='[公司A]', mode='exact'),
        WordReplaceRule(find='张三', replace='李四', mode='exact')
    ]
    re_engine = RuleEngine(terms=['13800138000'])
    pipeline = WordPipeline(rule_engine=re_engine, rules=rules)
    
    res = pipeline.process_document(doc_path, out_path)
    assert res['total_replacements'] >= 3
    
    # Verify resulting document
    doc_res = Document(str(out_path))
    full_text = doc_res.paragraphs[0].text
    assert '[公司A]' in full_text
    assert '李四' in full_text
    assert '北京科技有限公司' not in full_text
    assert '张三' not in full_text
    
    table_text = doc_res.tables[0].cell(0, 0).paragraphs[0].text
    assert '13800138000' not in table_text
    assert '*' * 11 in table_text
    
    # Verify format was preserved on inserted runs
    runs = doc_res.paragraphs[0].runs
    assert any(r.bold for r in runs if '[公司A]' in r.text)


def test_word_pipeline_document_rules_passport_and_phone_scan():
    doc = Document()
    doc.add_paragraph('护照号码 E12345678，手机 13800138000')

    pipeline = WordPipeline.from_document_rules()
    scan = pipeline.scan_document(doc)

    assert scan['total_matches'] >= 2
    matches = scan['paragraphs'][0]['matches']
    matched_text = {m['text'] for m in matches}
    assert 'E12345678' in matched_text
    assert '13800138000' in matched_text


def test_word_pipeline_passport_two_letter_prefix():
    """港澳/电子护照等 EL9115580 格式（DocumentRules 通用 1-2 字母 + 数字）。"""
    doc = Document()
    doc.add_paragraph('本人护照号码：EL9115580，赴日签证申请')

    pipeline = WordPipeline.from_document_rules()
    scan = pipeline.scan_document(doc)

    assert scan['total_matches'] >= 1
    matched_text = {m['text'] for m in scan['paragraphs'][0]['matches']}
    assert 'EL9115580' in matched_text


def test_word_pipeline_hk_phone_and_asterisk_replacement():
    doc = Document()
    doc.add_paragraph('联系电话：+852 84957302')

    pipeline = WordPipeline.from_document_rules()
    scan = pipeline.scan_document(doc)
    matches = scan['paragraphs'][0]['matches']
    assert any(m['text'] == '+852 84957302' for m in matches)
    hk_match = next(m for m in matches if m['text'] == '+852 84957302')
    assert hk_match['replacement'] == '*' * len('+852 84957302')


def test_word_pipeline_document_rules_redact_uses_stars(tmp_path):
    doc_path = tmp_path / 'mask.docx'
    out_path = tmp_path / 'mask_out.docx'
    doc = Document()
    doc.add_paragraph('护照 EL9115580')
    doc.save(str(doc_path))

    pipeline = WordPipeline.from_document_rules()
    pipeline.process_document(doc_path, out_path)
    result_text = Document(str(out_path)).paragraphs[0].text
    assert 'EL9115580' not in result_text
    assert '*' * 9 in result_text


def test_load_all_word_detection_rules_includes_saved_pii(tmp_path, monkeypatch):
    from core.word.rules_loader import load_all_word_detection_rules

    rules_file = tmp_path / 'pii_rules.json'
    rules_file.write_text(
        '''
        {
          "pii_rules": {
            "rule_custom": {
              "name": "内部代号",
              "category": "自定义",
              "pattern": "SECRET-\\\\d{4}",
              "enabled": true
            }
          },
          "word_replace_rules": []
        }
        ''',
        encoding='utf-8',
    )

    merged = load_all_word_detection_rules(path=rules_file)
    finds = [rule.find for rule in merged]
    assert any('SECRET' in find for find in finds)
