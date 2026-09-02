# -*- coding: utf-8 -*-
"""Generate synthetic test samples for ZonKey browser UI testing.

All sensitive terms are generic (CONFIDENTIAL / PROPRIETARY / INTERNAL USE ONLY)
plus format-valid fake PII. No vendor-specific words (public release rule).

Phase M note: PDF samples are built with reportlab (BSD) — no PyMuPDF/AGPL.
Sample coordinates keep the display convention of the previous fitz version
(origin top-left, y down) and are converted to reportlab's user space here.
"""
from pathlib import Path

from docx import Document
from reportlab.pdfgen import canvas

OUT = Path(__file__).parent


def _rl_y(h: float, y: float) -> float:
    """显示空间 y（向下）→ reportlab 用户空间 y（向上）。"""
    return h - y


# ---------------------------------------------------------------- drawing PDF
def make_drawing():
    W, H = 1191, 842  # A3 landscape
    path = OUT / "sample_drawing.pdf"
    c = canvas.Canvas(str(path), pagesize=(W, H))
    blue = (0, 0, 0.55)
    black = (0, 0, 0)
    grey = (0.75, 0.75, 0.75)
    darkred = (0.6, 0, 0)
    thin = 0.7

    def rect(x0, y0, x1, y1, color=blue, width=1.2):
        c.setStrokeColorRGB(*color)
        c.setLineWidth(width)
        c.rect(x0, _rl_y(H, y1), x1 - x0, y1 - y0, stroke=1, fill=0)

    def line(x0, y0, x1, y1, color=blue, width=thin):
        c.setStrokeColorRGB(*color)
        c.setLineWidth(width)
        c.line(x0, _rl_y(H, y0), x1, _rl_y(H, y1))

    def text(x, y, s, size=9, color=blue, rotate=None):
        c.setFillColorRGB(*color)
        if rotate:
            c.saveState()
            c.translate(x, _rl_y(H, y))
            c.rotate(rotate)
            c.drawString(0, 0, s)
            c.restoreState()
        else:
            c.drawString(x, _rl_y(H, y), s)

    # double frame border
    rect(15, 15, W - 15, H - 15, width=1.4)
    rect(25, 25, W - 25, H - 25, width=thin)

    # grid reference letters along frame
    for i, ch in enumerate("ABCDEFGH"):
        x = 60 + i * 135
        text(x, 22, ch, size=7)
        text(x, H - 19, ch, size=7)

    # ---- geometry: flange plate ----
    rect(160, 150, 560, 470)
    c.setStrokeColorRGB(*blue)
    c.setLineWidth(1.2)
    c.circle(360, _rl_y(H, 310), 75, stroke=1, fill=0)
    c.circle(360, _rl_y(H, 310), 22, stroke=1, fill=0)
    for cx, cy in [(215, 195), (505, 195), (215, 425), (505, 425)]:
        c.setStrokeColorRGB(*blue)
        c.setLineWidth(thin)
        c.circle(cx, _rl_y(H, cy), 18, stroke=1, fill=0)
        line(cx, cy - 40, cx, cy + 40)
        line(cx - 40, cy, cx + 40, cy)

    # dimension lines
    line(160, 500, 560, 500)
    text(330, 496, "400", size=9)
    line(120, 150, 120, 470)
    text(112, 315, "320", size=9, rotate=90)

    # ---- title block (bottom-right, boxed cells -> box detection target) ----
    tx0, ty0, tx1, ty1 = 700, 690, W - 25, H - 25
    rect(tx0, ty0, tx1, ty1)
    row_h = (ty1 - ty0) / 4
    col1 = tx0 + 250
    for i in range(1, 4):
        line(tx0, ty0 + i * row_h, tx1, ty0 + i * row_h)
    line(col1, ty0, col1, ty1)

    def cell(row, col, label, value):
        y = ty0 + row * row_h
        text(tx0 + 8 if col == 0 else col1 + 8, y + row_h - 14, label, size=7)
        text(tx0 + 60 if col == 0 else col1 + 80, y + row_h - 16, value, size=10, color=black)

    cell(0, 0, "TITLE", "FLANGE PLATE D3")
    cell(0, 1, "DWG NO", "ZS-2026-088")
    cell(1, 0, "MATERIAL", "SS316L")
    cell(1, 1, "REV", "B")
    cell(2, 0, "DRAWN BY", "Z. LI 2026-08-29")
    cell(2, 1, "SCALE", "1:2")
    cell(3, 0, "MARKING", "CONFIDENTIAL")
    cell(3, 1, "SHEET", "1 OF 1")

    # ---- notes block (boxed) ----
    nx0, ny0, nx1, ny1 = 640, 60, W - 25, 300
    rect(nx0, ny0, nx1, ny1)
    text(nx0 + 10, ny0 + 22, "NOTES:", size=11)
    notes = [
        "1. THIS DRAWING CONTAINS PROPRIETARY INFORMATION.",
        "2. DO NOT COPY OR DISTRIBUTE WITHOUT APPROVAL.",
        "3. DIMENSIONS IN MILLIMETERS. TOLERANCE +/-0.2.",
        "4. SURFACE FINISH RA 1.6 UNLESS NOTED.",
        "5. PART NO. ZS-088-FL-316L-QTY 4 REQUIRED.",
        "6. INTERNAL USE ONLY - PROJECT CODE PX-42.",
    ]
    for i, t in enumerate(notes):
        text(nx0 + 15, ny0 + 48 + i * 34, t, size=9, color=black)
        line(nx0 + 8, ny0 + 36 + i * 34, nx1 - 8, ny0 + 36 + i * 34, color=grey)

    # ---- margin stamps inside boxed strip (top-left) ----
    rect(40, 40, 380, 90, width=thin)
    text(55, 62, "CONFIDENTIAL", size=13, color=darkred)
    text(55, 82, "DO NOT COPY - RESTRICTED", size=9, color=darkred)

    c.showPage()
    c.save()
    print("sample_drawing.pdf OK")


# ------------------------------------------------------------- official PDF
def make_doc_pdf():
    W, H = 595, 842  # A4
    path = OUT / "sample_doc.pdf"
    c = canvas.Canvas(str(path), pagesize=(W, H))
    black = (0, 0, 0)

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    if "MSyh" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("MSyh", r"C:\Windows\Fonts\msyh.ttc", subfontIndex=0))

    def text(x, y, s, size=11, color=black, font="MSyh"):
        c.setFillColorRGB(*color)
        c.setFont(font, size)
        c.drawString(x, _rl_y(H, y), s)

    # red-header style official doc
    text(175, 90, "内部工作通知", size=22, color=(0.7, 0, 0))
    c.setStrokeColorRGB(0.7, 0, 0)
    c.setLineWidth(2)
    c.line(70, _rl_y(H, 110), 525, _rl_y(H, 110))
    text(70, 135, "INTERNAL USE ONLY", size=10, color=(0.6, 0, 0), font="Helvetica")
    text(380, 135, "编号：ZN-2026-042", size=10)

    body = [
        "各部门负责人：",
        "为配合三季度信息安全专项检查，现将人员联络信息汇总如下，",
        "本文件含 CONFIDENTIAL 内容，请勿外传。",
        "",
        "联系人：张三    手机：13812345678",
        "证件号码：11010119900307889X",
        "电子邮箱：zhangsan@example.com",
        "办公电话：010-66889900 转 1234",
        "备用联络：+85291234567（香港办事处）",
        "",
        "采购事项：设备合同金额 ¥1250000.00，已付款 850000元。",
        "付款账户：6222021234567890",
        "供应商信用代码：91110000MA01A2B3X4",
        "随附护照信息：E12345678",
        "",
        "请各部门于 2026 年 9 月 5 日前完成自查并回执。",
        "RESTRICTED - 本文件阅后回收。",
    ]
    y = 185
    for t in body:
        if t:
            text(70, y, t, size=11)
        y += 26

    text(420, 790, "综合管理部", size=11)
    text(420, 810, "2026-08-29", size=11)

    c.showPage()
    c.save()
    print("sample_doc.pdf OK")


# ------------------------------------------------------------------- docx
def make_docx():
    doc = Document()
    doc.add_heading("三季度信息安全自查报告", level=1)
    doc.add_paragraph("CONFIDENTIAL — 仅供内部使用（INTERNAL USE ONLY）")
    doc.add_paragraph("编制部门：综合管理部    日期：2026-08-29")
    doc.add_paragraph("")
    doc.add_heading("一、人员联络清单", level=2)
    doc.add_paragraph("联系人：李四，手机 13998765432，邮箱 lisi@example.com。")
    doc.add_paragraph("证件号码：440301198811116672。")
    doc.add_heading("二、财务信息", level=2)
    doc.add_paragraph("咨询服务费合计 ¥98000.00，经办账户 6222028765432109。")
    doc.add_paragraph("供应商信用代码：91440300MA5DA1B2C3。")
    doc.add_heading("三、保密要求", level=2)
    doc.add_paragraph("本报告含 PROPRIETARY 信息：DO NOT DISTRIBUTE。")
    doc.add_paragraph("如有疑问请联系 +85291239876（港澳联络）。")
    doc.save(OUT / "sample.docx")
    print("sample.docx OK")


if __name__ == "__main__":
    make_drawing()
    make_doc_pdf()
    make_docx()
