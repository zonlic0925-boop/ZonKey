# -*- coding: utf-8 -*-
"""生成全功能回归测试夹具：PDF / DOCX / WAV / MP4 / PNG。"""
import os
import shutil
import subprocess
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from PIL import Image, ImageDraw

HERE = Path(__file__).parent
FIX = HERE / "fixtures"
FIX.mkdir(parents=True, exist_ok=True)


def find_ffmpeg():
    exe = shutil.which("ffmpeg")
    if exe:
        return exe
    winget_root = Path(os.environ.get("LOCALAPPDATA", "")) / "Microsoft" / "WinGet" / "Packages"
    for pattern in ("Gyan.FFmpeg*/ffmpeg-*/bin/ffmpeg.exe", "FFmpeg*/ffmpeg-*/bin/ffmpeg.exe"):
        hits = sorted(winget_root.glob(pattern))
        if hits:
            return str(hits[-1])
    raise RuntimeError("ffmpeg not found")


def make_drawing_pdf():
    """工程图纸样张：外框 + 框内敏感词（走 rules/sensitive_terms.txt 词表命中）。"""
    doc = fitz.open()
    page = doc.new_page(width=842, height=595)  # A4 横向
    page.draw_rect(fitz.Rect(40, 40, 802, 555), color=(0, 0, 0), width=1.5)
    page.draw_rect(fitz.Rect(60, 60, 500, 300), color=(0, 0, 0), width=1)
    page.insert_text((80, 100), "CONFIDENTIAL", fontsize=20, fontname="helv")
    page.insert_text((80, 140), "PROPRIETARY DESIGN - DO NOT COPY", fontsize=12, fontname="helv")
    page.insert_text((80, 200), "Assembly drawing A-100", fontsize=10, fontname="helv")
    page.insert_text((80, 520), "Sheet 1 of 3", fontsize=9, fontname="helv")
    doc.save(FIX / "drawing.pdf")
    doc.close()


def make_document_pdf():
    """公文样张：整行文字含敏感标记 + PII（手机号/邮箱，命中 pii_rules 正则）。"""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 90), "Quarterly Report", fontsize=18, fontname="hebo")
    page.insert_text((72, 130), "INTERNAL USE ONLY - RESTRICTED", fontsize=12, fontname="helv")
    page.insert_text((72, 160), "Revenue increased by 12 percent this quarter.", fontsize=11, fontname="helv")
    page.insert_text((72, 190), "Contact: 13800138000  zhangsan@example.com", fontsize=11, fontname="helv")
    doc.save(FIX / "document.pdf")
    doc.close()


def make_docx():
    d = Document()
    d.add_heading("Test Document", level=1)
    d.add_paragraph("This document is CONFIDENTIAL and should not be distributed.")
    d.add_paragraph("Normal paragraph without any secrets.")
    d.save(FIX / "sample.docx")


def make_speech_wav():
    """Windows SAPI TTS 合成真实语音，用于转写端到端验证。"""
    out = FIX / "speech.wav"
    ps = (
        "Add-Type -AssemblyName System.Speech; "
        "$s = New-Object System.Speech.Synthesis.SpeechSynthesizer; "
        f"$s.SetOutputToWaveFile('{out.as_posix()}'); "
        "$s.Speak('Hello world, this is a transcription engine test.'); "
        "$s.Dispose()"
    )
    subprocess.run(["powershell", "-NoProfile", "-Command", ps], check=True, timeout=60)
    assert out.exists() and out.stat().st_size > 10000


def make_clip_mp4():
    out = FIX / "clip.mp4"
    subprocess.run(
        [find_ffmpeg(), "-hide_banner", "-y",
         "-f", "lavfi", "-i", "testsrc=duration=2:size=320x240:rate=15",
         "-f", "lavfi", "-i", "sine=frequency=440:duration=2",
         "-shortest", "-pix_fmt", "yuv420p", str(out)],
        check=True, capture_output=True, timeout=120,
    )
    assert out.exists() and out.stat().st_size > 10000


def make_images():
    for name, color in (("img1.png", (200, 60, 60)), ("img2.png", (60, 120, 200))):
        img = Image.new("RGB", (320, 240), color)
        d = ImageDraw.Draw(img)
        d.rectangle([20, 20, 300, 220], outline=(0, 0, 0), width=3)
        img.save(FIX / name)


if __name__ == "__main__":
    make_drawing_pdf()
    make_document_pdf()
    make_docx()
    make_images()
    make_clip_mp4()
    make_speech_wav()
    for f in sorted(FIX.iterdir()):
        print(f"{f.name:16s} {f.stat().st_size/1024:8.1f} KB")
    print("FIXTURES OK")
