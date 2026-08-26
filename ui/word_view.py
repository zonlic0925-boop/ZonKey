from __future__ import annotations

import logging
from pathlib import Path
from PyQt5.QtCore import Qt
from PyQt5.QtWidgets import (
    QHBoxLayout,
    QLabel,
    QSplitter,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)

import docx

logger = logging.getLogger(__name__)


class WordCompareView(QWidget):
    """Word 文档双栏对照预览视口：左侧原文档流文本预览，右侧脱敏后文档文本对照预览。"""

    def __init__(self, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self._build_ui()

    def _build_ui(self) -> None:
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        splitter = QSplitter(Qt.Horizontal)

        # 左栏：原文档预览
        w_orig = QWidget()
        l_orig = QVBoxLayout(w_orig)
        l_orig.setContentsMargins(4, 4, 4, 4)
        self.lbl_orig = QLabel("📄 原文档内容 (只读预览)")
        self.txt_orig = QTextEdit()
        self.txt_orig.setReadOnly(True)
        l_orig.addWidget(self.lbl_orig)
        l_orig.addWidget(self.txt_orig)
        splitter.addWidget(w_orig)

        # 右栏：脱敏后预览
        w_redacted = QWidget()
        l_redacted = QVBoxLayout(w_redacted)
        l_redacted.setContentsMargins(4, 4, 4, 4)
        self.lbl_redacted = QLabel("🛡️ 脱敏后内容 (只读预览)")
        self.txt_redacted = QTextEdit()
        self.txt_redacted.setReadOnly(True)
        l_redacted.addWidget(self.lbl_redacted)
        l_redacted.addWidget(self.txt_redacted)
        splitter.addWidget(w_redacted)

        splitter.setSizes([500, 500])

        # 滚动条双向联动同步
        self.txt_orig.verticalScrollBar().valueChanged.connect(
            self.txt_redacted.verticalScrollBar().setValue
        )
        self.txt_redacted.verticalScrollBar().valueChanged.connect(
            self.txt_orig.verticalScrollBar().setValue
        )

        layout.addWidget(splitter)

    def load_document(self, orig_path: str | Path, redacted_path: str | Path | None = None) -> None:
        """加载原文档和（可选的）脱敏后 Word 文档文本。"""
        p_orig = Path(orig_path)
        if p_orig.is_file() and p_orig.suffix.lower() == ".docx":
            try:
                doc = docx.Document(p_orig)
                txt = "\n".join([p.text for p in doc.paragraphs])
                self.txt_orig.setPlainText(txt)
                self.lbl_orig.setText(f"📄 原文档: {p_orig.name}")
            except Exception as e:
                logger.warning("解析原 docx 失败: %s", e)
                self.txt_orig.setPlainText(f"(无法解析 Word 文档: {e})")
        else:
            self.txt_orig.clear()

        if redacted_path:
            p_red = Path(redacted_path)
            if p_red.is_file() and p_red.suffix.lower() == ".docx":
                try:
                    doc = docx.Document(p_red)
                    txt = "\n".join([p.text for p in doc.paragraphs])
                    self.txt_redacted.setPlainText(txt)
                    self.lbl_redacted.setText(f"🛡️ 脱敏后: {p_red.name}")
                except Exception as e:
                    logger.warning("解析脱敏 docx 失败: %s", e)
                    self.txt_redacted.setPlainText(f"(无法解析 Word 文档: {e})")
            else:
                self.txt_redacted.clear()
                self.lbl_redacted.setText("🛡️ 脱敏后内容 (尚未执行脱敏)")
        else:
            self.txt_redacted.clear()
            self.lbl_redacted.setText("🛡️ 脱敏后内容 (尚未执行脱敏)")

    def clear(self) -> None:
        self.txt_orig.clear()
        self.txt_redacted.clear()
        self.lbl_orig.setText("📄 原文档内容 (只读预览)")
        self.lbl_redacted.setText("🛡️ 脱敏后内容 (只读预览)")


WordView = WordCompareView
__all__ = ["WordCompareView", "WordView"]